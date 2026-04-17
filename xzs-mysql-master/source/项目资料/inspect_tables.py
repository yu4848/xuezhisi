from docx import Document
import os

base = os.path.dirname(os.path.abspath(__file__))

# ── 检查格式模板表格内容 ─────────────────────────────────────
print("=== 论文格式.docx 表格 ===")
fmt = Document(os.path.join(base, '论文格式.docx'))
for i, tbl in enumerate(fmt.tables):
    print(f"\n--- 表格 {i} ---")
    for row in tbl.rows:
        cells = [c.text.strip() for c in row.cells]
        print(" | ".join(cells))

# ── 检查当前论文内容.docx 的样式使用情况 ────────────────────
print("\n\n=== 论文内容.docx 前50段 样式+内容 ===")
content = Document(os.path.join(base, '论文内容.docx'))
for i, para in enumerate(content.paragraphs[:50]):
    if not para.text.strip():
        continue
    pf = para.paragraph_format
    run = para.runs[0] if para.runs else None
    print(f"[{para.style.name}] align={para.alignment} fi={pf.first_line_indent} ls={pf.line_spacing} "
          f"font={run.font.name if run else '-'} size={run.font.size if run else '-'} "
          f"| {para.text[:60]}")

# ── 列出论文内容.docx 全部样式 ─────────────────────────────
print("\n=== 论文内容.docx 样式列表 ===")
for s in content.styles:
    if s.type.name == 'PARAGRAPH':
        print(f"  {s.name}")
