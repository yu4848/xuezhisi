"""
apply_format.py
按广东技术师范大学本科毕业论文规范（理工类）格式化 论文内容.docx
输出: 论文内容_格式化.docx（原文件保留不变）

格式依据：论文格式.docx 中 表6.1 正文格式（理工外语类）
┌──────┬──────────────┬────────────────┬───────────────────┐
│ 层级 │    字体      │      行距      │      对齐         │
├──────┼──────────────┼────────────────┼───────────────────┤
│ 第1层│三号黑体加粗  │ 段前段后1行    │ 顶格              │
│ 第2层│小三黑体加粗  │ 段前段后0.5行  │ 顶格              │
│ 第3层│四号黑体加粗  │ 段前段后0.5行  │ 顶格              │
│ 第4层│小四黑体加粗  │ 1.5倍行距      │ 首行缩进2字       │
│ 正文 │小四宋体      │ 1.5倍行距      │ 首行缩进2字       │
└──────┴──────────────┴────────────────┴───────────────────┘

页面设置：A4，上2.8  下2.2  左3.0  右2.0（cm），页眉1.8  页脚1.4
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.abspath(__file__))
SRC  = os.path.join(BASE, '论文内容.docx')
OUT  = os.path.join(BASE, '论文内容_格式化.docx')

# ── 字号常量（磅）──────────────────────────────────────────
PT_ER    = 22   # 二号
PT_SAN   = 16   # 三号
PT_XISAN = 15   # 小三
PT_SI    = 14   # 四号
PT_XISI  = 12   # 小四
PT_WU    = 10.5 # 五号
PT_XIWU  = 9    # 小五

J   = WD_ALIGN_PARAGRAPH.JUSTIFY
C   = WD_ALIGN_PARAGRAPH.CENTER
L   = WD_ALIGN_PARAGRAPH.LEFT
OPF = WD_LINE_SPACING.ONE_POINT_FIVE
SGL = WD_LINE_SPACING.SINGLE


# ═══════════════════════════════════════════════════════════
# 辅助函数
# ═══════════════════════════════════════════════════════════

def _get_or_create(parent, tag):
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el


def set_style_fonts(style, cn=None, latin=None, size=None, bold=None, color_rgb=None):
    """设置样式的字体属性（含中文 eastAsia 字体）"""
    rPr = style.font._element          # w:rPr 节点（不存在会自动创建）

    # ── rFonts ──────────────────────────────────────────────
    if cn or latin:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        if latin:
            rFonts.set(qn('w:ascii'),  latin)
            rFonts.set(qn('w:hAnsi'),  latin)
        if cn:
            rFonts.set(qn('w:eastAsia'), cn)
            rFonts.set(qn('w:cs'),       cn)

    # ── 字号（半角点：1pt = 2 半角点）──────────────────────
    if size is not None:
        hp = str(int(size * 2))
        _get_or_create(rPr, 'w:sz').set(qn('w:val'), hp)
        _get_or_create(rPr, 'w:szCs').set(qn('w:val'), hp)

    # ── 加粗 ────────────────────────────────────────────────
    if bold is not None:
        style.font.bold = bold

    # ── 颜色 ────────────────────────────────────────────────
    if color_rgb is not None:
        style.font.color.rgb = color_rgb


def set_style_para(style, align=None, sb=None, sa=None,
                   ls=None, lsr=None, fi=None, li=None):
    """设置样式的段落格式"""
    pf = style.paragraph_format
    if align is not None: pf.alignment         = align
    if sb    is not None: pf.space_before       = sb
    if sa    is not None: pf.space_after        = sa
    if ls    is not None: pf.line_spacing       = ls
    if lsr   is not None: pf.line_spacing_rule  = lsr
    if fi    is not None: pf.first_line_indent  = fi
    if li    is not None: pf.left_indent        = li


BLACK = RGBColor(0, 0, 0)

# ── 首行缩进：小四2字 = 2×12pt = 24pt ──────────────────────
BODY_FI = Pt(24)

doc = Document(SRC)


# ═══════════════════════════════════════════════════════════
# 1. 页面设置
# ═══════════════════════════════════════════════════════════
for sec in doc.sections:
    sec.top_margin      = Cm(2.8)
    sec.bottom_margin   = Cm(2.2)
    sec.left_margin     = Cm(3.0)
    sec.right_margin    = Cm(2.0)
    sec.header_distance = Cm(1.8)
    sec.footer_distance = Cm(1.4)
print('✓ 页面设置完成')


# ═══════════════════════════════════════════════════════════
# 2. Heading 1 — 论文标题（二号黑体加粗，居中，段前后1行）
#    Pandoc 把 # 标题映射为 Heading 1
# ═══════════════════════════════════════════════════════════
h1 = doc.styles['Heading 1']
set_style_para(h1, align=C, sb=Pt(PT_ER), sa=Pt(PT_ER),
               ls=1.5, lsr=OPF, fi=Pt(0), li=Pt(0))
set_style_fonts(h1, cn='黑体', latin='Times New Roman',
                size=PT_ER, bold=True, color_rgb=BLACK)
print('✓ Heading 1 完成')


# ═══════════════════════════════════════════════════════════
# 3. Heading 2 — 第一层（三号黑体加粗，顶格，段前后1行）
#    Pandoc 把 ## 映射为 Heading 2（第X章 / 摘要 / 附录等）
# ═══════════════════════════════════════════════════════════
h2 = doc.styles['Heading 2']
set_style_para(h2, align=L, sb=Pt(PT_SAN * 1.5), sa=Pt(PT_SAN * 1.5),
               ls=1.5, lsr=OPF, fi=Pt(0), li=Pt(0))
set_style_fonts(h2, cn='黑体', latin='Times New Roman',
                size=PT_SAN, bold=True, color_rgb=BLACK)
print('✓ Heading 2 完成')


# ═══════════════════════════════════════════════════════════
# 4. Heading 3 — 第二层（小三黑体加粗，顶格，段前后0.5行）
#    Pandoc 把 ### 映射为 Heading 3（X.X节）
# ═══════════════════════════════════════════════════════════
h3 = doc.styles['Heading 3']
set_style_para(h3, align=L, sb=Pt(PT_XISAN * 0.75), sa=Pt(PT_XISAN * 0.75),
               ls=1.5, lsr=OPF, fi=Pt(0), li=Pt(0))
set_style_fonts(h3, cn='黑体', latin='Times New Roman',
                size=PT_XISAN, bold=True, color_rgb=BLACK)
print('✓ Heading 3 完成')


# ═══════════════════════════════════════════════════════════
# 5. Heading 4 — 第三层（四号黑体加粗，顶格，段前后0.5行）
#    Pandoc 把 #### 映射为 Heading 4（X.X.X条）
# ═══════════════════════════════════════════════════════════
h4 = doc.styles['Heading 4']
set_style_para(h4, align=L, sb=Pt(PT_SI * 0.75), sa=Pt(PT_SI * 0.75),
               ls=1.5, lsr=OPF, fi=Pt(0), li=Pt(0))
set_style_fonts(h4, cn='黑体', latin='Times New Roman',
                size=PT_SI, bold=True, color_rgb=BLACK)
print('✓ Heading 4 完成')


# ═══════════════════════════════════════════════════════════
# 6. 正文样式（小四宋体，1.5倍，首行缩进2字，两端对齐）
# ═══════════════════════════════════════════════════════════
BODY_STYLES = ('Normal', 'Body Text', 'First Paragraph', 'Compact', 'Abstract')
for sname in BODY_STYLES:
    try:
        s = doc.styles[sname]
        set_style_para(s, align=J, sb=Pt(0), sa=Pt(0),
                       ls=1.5, lsr=OPF, fi=BODY_FI, li=Pt(0))
        set_style_fonts(s, cn='宋体', latin='Times New Roman',
                        size=PT_XISI, bold=False, color_rgb=BLACK)
        print(f'✓ {sname} 完成')
    except KeyError:
        pass


# ═══════════════════════════════════════════════════════════
# 7. 代码块（Courier New 小五9pt，顶格，单倍行距，段前后3pt）
# ═══════════════════════════════════════════════════════════
try:
    sc = doc.styles['Source Code']
    set_style_para(sc, align=L, sb=Pt(3), sa=Pt(3),
                   ls=1.0, lsr=SGL, fi=Pt(0), li=Pt(0))
    set_style_fonts(sc, cn='Courier New', latin='Courier New',
                    size=PT_XIWU, bold=False, color_rgb=BLACK)
    print('✓ Source Code 完成')
except KeyError:
    print('  Source Code 样式不存在，跳过')


# ═══════════════════════════════════════════════════════════
# 8. 表格内文字（五号宋体）
# ═══════════════════════════════════════════════════════════
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                pf = para.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after  = Pt(0)
                if pf.line_spacing is None:
                    pf.line_spacing      = 1.0
                    pf.line_spacing_rule = SGL
print('✓ 表格段落行距完成')


# ═══════════════════════════════════════════════════════════
# 9. 保存
# ═══════════════════════════════════════════════════════════
doc.save(OUT)
print(f'\n✓✓ 已保存至: {OUT}')
